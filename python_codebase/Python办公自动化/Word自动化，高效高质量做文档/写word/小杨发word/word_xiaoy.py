# -*- coding: utf-8 -*-
# All Rights Reserved 
# @Time    : 2022/10/914:31
# @Author  : Small Fu
# @Email   : 2737454073@qq.com
# @File    : word_xiaoy.py
__author__ = 'Small Fu'

from docx import Document  # 初始化word
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 段落对其
from docx.shared import Pt  # 字体磅数
from docx.oxml.ns import qn  # 设置中文格式

import time

price = input("今日价格：")
# 这是客户
customer_list = [f"客户{i}" for i in range(10)]

time_sameDay = time.strftime("%Y年 %m月 %d日", time.localtime())

# 开始设置
for customer in customer_list:
    doc = Document()
    # 设置基础字体,这个值设置英文和数字
    doc.styles["Normal"].font.name = u"宋体"
    # 设置基础字符大小
    doc.styles["Normal"].font.size = Pt(14)
    # 增加中文支持
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")
    # 初始化一个段落
    p1 = doc.add_paragraph()
    # 设置文字对齐居中，没有的话默认左对齐
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 添加第一段的内容
    run1 = p1.add_run(f"关于下达{time_sameDay}产品价格通知")
    # 设置这一段的字体是微软雅黑西文
    run1.font.name = "微软雅黑"
    # 设置中文
    run1._element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")
    # 设置字体大小
    run1.font.size = Pt(21)
    # 是否加粗，默认不加粗，False
    run1.font.bold = True
    # 设置这一段前面的字体大小和后面的
    p1.space_after = Pt(5)
    p1.space_before = Pt(5)

    # 第二段内容
    p2 = doc.add_paragraph()
    # 这里是添加尊称
    run2 = p2.add_run(f"{customer}: \n  您好")
    # 设置这一段的字体是微软雅黑西文
    run2.font.name = "仿宋_GB2312"
    # 设置中文
    run2._element.rPr.rFonts.set(qn("w:eastAsia"), u"仿宋_GB2312 ")
    # 设置字体大小
    run2.font.size = Pt(16)
    # 是否加粗，默认不加粗，False
    run2.font.bold = True

    # 第三段内容
    p3 = doc.add_paragraph()
    # 这里是添加尊称
    run3 = p3.add_run(f"    根据公司安排，为提供优质客户需求，我单位拟定了今日黄金价格为{price}元，特此通知")
    # 设置这一段的字体是微软雅黑西文
    run3.font.name = "仿宋_GB2312"
    # 设置中文
    run3._element.rPr.rFonts.set(qn("w:eastAsia"), u"仿宋_GB2312 ")
    # 设置字体大小
    run3.font.size = Pt(16)
    # 是否加粗，默认不加粗，False
    run3.font.bold = True

    # 第四段内容
    p4 = doc.add_paragraph()
    # 这里是添加尊称
    run4 = p4.add_run("     联系人：小符    电话：176XXXXXX44    ")
    # 设置这一段的字体是微软雅黑西文
    run4.font.name = "仿宋_GB2312"
    # 设置中文
    run4._element.rPr.rFonts.set(qn("w:eastAsia"), u"仿宋_GB2312 ")
    # 设置字体大小
    run4.font.size = Pt(16)
    # 是否加粗，默认不加粗，False
    run4.font.bold = True

    # 保存
    doc.save(f"{customer}-价格通知.docx")
