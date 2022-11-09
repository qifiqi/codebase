# -*- coding: utf-8 -*-
# All Rights Reserved 
# @Time    : 2022/10/916:23
# @Author  : Small Fu
# @Email   : 2737454073@qq.com
# @File    : 函数式模板.py
__author__ = 'Small Fu'

from docx import Document  # 初始化word
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 段落对其
from docx.shared import Pt  # 字体磅数
from docx.oxml.ns import qn  # 设置中文格式

doc = Document()
doc.styles["Normal"].font.name = u"黑体"
doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), u"黑体")


def add_context(context):
    p1 = doc.add_paragraph()
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 添加内容
    run1 = p1.add_run(str(context))
    # 设置这一段的字体是微软雅黑西文
    run1.font.size = Pt(16)
    # 设置这一段前面的字体大小和后面的
    p1.space_after = Pt(5)
    p1.space_before = Pt(5)
