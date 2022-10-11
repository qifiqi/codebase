import docx

fn = "./HCIP.docx"
doc = docx.Document(fn)

aa = []

for paragraph in doc.paragraphs:
    aa.append(paragraph.text)

for i in aa:
    print(i)
    print('*'*88)

# for table in doc.tables:
#     for row in table.rows:
#         for cell in row.cells:
#             print("*"*88)
#             print(cell.text)
